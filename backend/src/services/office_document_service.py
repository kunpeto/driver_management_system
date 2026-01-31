"""
OfficeDocumentService Office 文件生成服務
對應 tasks.md T137: 實作 Office 文件生成服務
對應 spec.md: User Story 8 - Office 文件生成（後端實作）

使用 python-docx 生成 Word 文件，模板填充，條碼嵌入。
改為後端直接生成，返回二進位流。
"""

import io
import os
import re
from datetime import date, time
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Inches

from src.models import (
    AssessmentNotice,
    AssessmentType,
    CorrectiveMeasures,
    Employee,
    EventInvestigation,
    PersonnelInterview,
    Profile,
    ProfileType,
)
from src.services.barcode_service import BarcodeService
from src.utils.barcode_encoder import BarcodeEncoder
from src.utils.file_naming import FileNaming


class OfficeDocumentServiceError(Exception):
    """Office 文件服務錯誤"""
    pass


class TemplateNotFoundError(OfficeDocumentServiceError):
    """模板不存在錯誤"""
    pass


class InvalidProfileTypeError(OfficeDocumentServiceError):
    """無效的履歷類型錯誤"""
    pass


class OfficeDocumentService:
    """
    Office 文件生成服務

    使用 python-docx 生成 Word 文件，支援模板填充和條碼嵌入。
    """

    # 模板目錄
    TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

    # 模板檔案對應
    TEMPLATE_FILES = {
        ProfileType.EVENT_INVESTIGATION.value: "event_investigation.docx",
        ProfileType.PERSONNEL_INTERVIEW.value: "personnel_interview.docx",
        ProfileType.CORRECTIVE_MEASURES.value: "corrective_measures.docx",
        # 考核通知根據類型選擇
    }

    # 考核通知模板
    ASSESSMENT_TEMPLATES = {
        AssessmentType.PLUS.value: "assessment_notice_plus.docx",
        AssessmentType.MINUS.value: "assessment_notice_minus.docx",
    }

    def __init__(self):
        """初始化服務"""
        self.barcode_service = BarcodeService()

    def generate(
        self,
        profile: Profile,
        employee: Employee,
        include_barcode: bool = True
    ) -> bytes:
        """
        生成 Office 文件

        Args:
            profile: 履歷物件（需含子表資料）
            employee: 員工物件
            include_barcode: 是否嵌入條碼

        Returns:
            Word 文件的 bytes

        Raises:
            InvalidProfileTypeError: 基本履歷不能生成文件
            TemplateNotFoundError: 模板不存在
        """
        if profile.profile_type == ProfileType.BASIC.value:
            raise InvalidProfileTypeError("基本履歷不能生成文件")

        # 取得模板路徑
        template_path = self._get_template_path(profile)

        # 載入模板
        doc = Document(template_path)

        # 準備佔位符資料
        placeholders = self._prepare_placeholders(profile, employee)

        # 替換佔位符
        self._replace_placeholders(doc, placeholders)

        # 嵌入條碼
        if include_barcode:
            self._embed_barcode(doc, profile)

        # 輸出到記憶體
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    def generate_filename(
        self,
        profile: Profile,
        employee: Employee,
        with_version: bool = False
    ) -> str:
        """
        生成檔案名稱

        Args:
            profile: 履歷物件
            employee: 員工物件
            with_version: 是否包含版本號

        Returns:
            檔案名稱
        """
        assessment_type = None
        if profile.assessment_notice:
            assessment_type = profile.assessment_notice.assessment_type

        if with_version:
            return FileNaming.generate_with_version(
                profile_type=profile.profile_type,
                event_date=profile.event_date,
                employee_name=employee.employee_name,
                version=profile.document_version,
                train_number=profile.train_number,
                event_location=profile.event_location,
                assessment_type=assessment_type,
            )
        else:
            return FileNaming.generate(
                profile_type=profile.profile_type,
                event_date=profile.event_date,
                employee_name=employee.employee_name,
                train_number=profile.train_number,
                event_location=profile.event_location,
                assessment_type=assessment_type,
            )

    def _get_template_path(self, profile: Profile) -> Path:
        """取得模板路徑"""
        if profile.profile_type == ProfileType.ASSESSMENT_NOTICE.value:
            # 考核通知根據類型選擇模板
            if profile.assessment_notice:
                assessment_type = profile.assessment_notice.assessment_type
                template_file = self.ASSESSMENT_TEMPLATES.get(
                    assessment_type,
                    "assessment_notice_minus.docx"
                )
            else:
                template_file = "assessment_notice_minus.docx"
        else:
            template_file = self.TEMPLATE_FILES.get(profile.profile_type)

        if not template_file:
            raise TemplateNotFoundError(
                f"找不到類型 {profile.profile_type} 的模板"
            )

        template_path = self.TEMPLATES_DIR / template_file

        if not template_path.exists():
            raise TemplateNotFoundError(f"模板檔案不存在: {template_path}")

        return template_path

    def _prepare_placeholders(
        self,
        profile: Profile,
        employee: Employee
    ) -> dict[str, Any]:
        """
        準備佔位符資料

        根據模板 README.md 的欄位映射表準備資料。
        """
        # 基本資料
        placeholders = {
            # 員工資料
            "employee_name": employee.employee_name,
            "employee_id": employee.employee_id,
            "hire_date": self._format_date(employee.hire_year_month),

            # 事件資料（Profile 主表）
            "event_date": self._format_date(profile.event_date),
            "event_time": self._format_time(profile.event_time),
            "event_location": profile.event_location or "",
            "train_number": profile.train_number or "",
            "event_title": profile.event_title or "",
            "event_description": profile.event_description or "",
            "data_source": profile.data_source or "",

            # 考核資料（Profile 主表共用）
            "assessment_item": profile.assessment_item or "",
            "assessment_score": str(profile.assessment_score) if profile.assessment_score else "",

            # 條碼
            "barcode_id": self._generate_barcode_id(profile),
        }

        # 根據類型添加子表資料
        if profile.profile_type == ProfileType.EVENT_INVESTIGATION.value:
            placeholders.update(self._prepare_event_investigation(profile))

        elif profile.profile_type == ProfileType.PERSONNEL_INTERVIEW.value:
            placeholders.update(self._prepare_personnel_interview(profile))

        elif profile.profile_type == ProfileType.CORRECTIVE_MEASURES.value:
            placeholders.update(self._prepare_corrective_measures(profile))

        elif profile.profile_type == ProfileType.ASSESSMENT_NOTICE.value:
            placeholders.update(self._prepare_assessment_notice(profile))

        return placeholders

    def _prepare_event_investigation(
        self,
        profile: Profile
    ) -> dict[str, Any]:
        """準備事件調查子表資料"""
        ei = profile.event_investigation
        if not ei:
            return {}

        return {
            "incident_cause": ei.cause_analysis or "",
            "incident_process": ei.process_description or "",
            "improvement_suggestion": ei.improvement_suggestions or "",
            "investigator": ei.investigator or "",
            "investigation_date": self._format_date(ei.investigation_date),
        }

    def _prepare_personnel_interview(
        self,
        profile: Profile
    ) -> dict[str, Any]:
        """準備人員訪談子表資料"""
        pi = profile.personnel_interview
        if not pi:
            return {}

        placeholders = {
            "shift1": pi.shift_event_day or "",
            "shift2": pi.shift_before_1day or "",
            "shift3": pi.shift_before_2days or "",
            "interviewer": pi.interviewer or "",
            "interview_date": self._format_date(pi.interview_date),
            "interview_content": pi.interview_content or "",
            "conclusion": pi.conclusion or "",
        }

        # 訪談結果勾選
        ir_data = pi.interview_result_data or {}
        for i in range(1, 8):
            key = f"ir_{i}"
            placeholders[key] = self._format_checkbox(ir_data.get(key, False))
        placeholders["ir_other_text"] = ir_data.get("ir_other_text", "")

        # 後續行動勾選
        fa_data = pi.follow_up_action_data or {}
        for i in range(1, 8):
            key = f"fa_{i}"
            placeholders[key] = self._format_checkbox(fa_data.get(key, False))
        placeholders["fa_other_text"] = fa_data.get("fa_other_text", "")

        return placeholders

    def _prepare_corrective_measures(
        self,
        profile: Profile
    ) -> dict[str, Any]:
        """準備矯正措施子表資料"""
        cm = profile.corrective_measures
        if not cm:
            return {}

        return {
            "incident_cause": cm.event_summary or "",
            "root_cause_analysis": cm.event_summary or "",
            "corrective_action": cm.corrective_actions or "",
        }

    def _prepare_assessment_notice(
        self,
        profile: Profile
    ) -> dict[str, Any]:
        """準備考核通知子表資料"""
        an = profile.assessment_notice
        if not an:
            return {}

        return {
            "assessment_type": an.assessment_type,
            "issue_date": self._format_date(an.issue_date),
            "approver": an.approver or "",
        }

    def _replace_placeholders(
        self,
        doc: Document,
        placeholders: dict[str, Any]
    ) -> None:
        """
        替換文件中的佔位符

        佔位符格式：{variable_name}
        """
        # 替換段落中的佔位符
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, placeholders)

        # 替換表格中的佔位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, placeholders)

    def _replace_in_paragraph(
        self,
        paragraph,
        placeholders: dict[str, Any]
    ) -> None:
        """替換段落中的佔位符"""
        if not paragraph.text:
            return

        # 使用正則表達式找到所有佔位符
        pattern = r'\{(\w+)\}'
        text = paragraph.text

        for match in re.finditer(pattern, text):
            key = match.group(1)
            if key in placeholders:
                value = str(placeholders[key])
                text = text.replace(f"{{{key}}}", value)

        # 只有在有變更時才更新
        if text != paragraph.text:
            paragraph.text = text

    def _embed_barcode(self, doc: Document, profile: Profile) -> None:
        """
        嵌入條碼到文件

        尋找包含 {barcode_image} 佔位符的段落，替換為條碼圖片。
        """
        barcode_placeholder = "{barcode_image}"

        # 生成條碼資料
        assessment_type = None
        if profile.assessment_notice:
            assessment_type = profile.assessment_notice.assessment_type

        barcode_data = BarcodeEncoder.encode(
            profile_id=profile.id,
            profile_type=profile.profile_type,
            version=profile.document_version,
            assessment_type=assessment_type
        )

        # 生成條碼圖片
        barcode_bytes = self.barcode_service.generate(
            barcode_data,
            height=10.0,  # 較小的條碼以適合文件
            include_text=True
        )

        # 在文件中尋找並替換
        for paragraph in doc.paragraphs:
            if barcode_placeholder in paragraph.text:
                # 清除段落文字
                paragraph.text = ""

                # 插入圖片
                run = paragraph.add_run()
                barcode_stream = io.BytesIO(barcode_bytes)
                run.add_picture(barcode_stream, width=Inches(2.5))
                break

        # 也檢查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if barcode_placeholder in paragraph.text:
                            paragraph.text = ""
                            run = paragraph.add_run()
                            barcode_stream = io.BytesIO(barcode_bytes)
                            run.add_picture(barcode_stream, width=Inches(2.0))
                            return  # 只嵌入一次

    def _generate_barcode_id(self, profile: Profile) -> str:
        """生成條碼 ID 字串"""
        assessment_type = None
        if profile.assessment_notice:
            assessment_type = profile.assessment_notice.assessment_type

        return BarcodeEncoder.encode(
            profile_id=profile.id,
            profile_type=profile.profile_type,
            version=profile.document_version,
            assessment_type=assessment_type
        )

    def _format_date(self, d: Optional[date]) -> str:
        """格式化日期"""
        if not d:
            return ""
        if isinstance(d, str):
            return d
        return d.strftime("%Y-%m-%d")

    def _format_time(self, t: Optional[time]) -> str:
        """格式化時間"""
        if not t:
            return ""
        if isinstance(t, str):
            return t
        return t.strftime("%H:%M")

    def _format_checkbox(self, checked: bool) -> str:
        """格式化勾選框"""
        return "V" if checked else ""
